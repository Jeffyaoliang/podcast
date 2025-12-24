# -*- coding: utf-8 -*-
"""
将Markdown文件转换为DOCX文档
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

def parse_markdown_to_docx(md_file, docx_file):
    """将Markdown文件转换为DOCX文档"""
    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            heading.style.font.name = '微软雅黑'
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=2)
            heading.style.font.name = '微软雅黑'
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
            heading.style.font.name = '微软雅黑'
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:].strip(), level=4)
            heading.style.font.name = '微软雅黑'
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
        elif line.strip().startswith('|'):
            # 表格行
            cells = [c.strip() for c in line.strip('|').split('|')]
            if len(cells) > 1:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for j, cell in enumerate(cells):
                    hdr_cells[j].text = cell
        elif line.strip().startswith('> '):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(text)
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
        elif line.strip() and not line.strip().startswith('---'):
            text = line.strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            parts = re.split(r'(\*\*.*?\*\*)', text)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                else:
                    p.add_run(part)
        i += 1
    
    doc.save(docx_file)
    print(f'成功将 {md_file} 转换为 {docx_file}')

if __name__ == '__main__':
    md_file = 'M2.1评测文章_多语言方向_优化版.md'
    docx_file = 'M2.1评测文章_多语言方向_完整版.docx'
    parse_markdown_to_docx(md_file, docx_file)

