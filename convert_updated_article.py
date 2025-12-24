#!/usr/bin/env python3
"""
Markdown转DOCX转换脚本
将M2.1评测文章转换为Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown
import sys
import os

def markdown_to_docx(md_file, docx_file):
    """将Markdown文件转换为DOCX"""
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 解析Markdown
    html_content = markdown.markdown(md_content, extensions=['extra', 'fenced_code', 'tables'])
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # 处理HTML内容并添加到文档
    lines = html_content.split('\n')
    
    in_code_block = False
    code_content = []
    
    for line in lines:
        line = line.strip()
        
        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_content:
                    p = doc.add_paragraph()
                    p.style = 'Code'
                    for code_line in code_content:
                        p.add_run(code_line + '\n')
                    code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # 处理HTML标签
        line = line.replace('<h1>', '').replace('</h1>', '')
        line = line.replace('<h2>', '').replace('</h2>', '')
        line = line.replace('<h3>', '').replace('</h3>', '')
        line = line.replace('<strong>', '').replace('</strong>', '')
        line = line.replace('<em>', '').replace('</em>', '')
        line = line.replace('<p>', '').replace('</p>', '')
        line = line.replace('<br>', '\n')
        line = line.replace('<ul>', '').replace('</ul>', '')
        line = line.replace('<li>', '• ').replace('</li>', '')
        line = line.replace('&gt;', '>')
        line = line.replace('&lt;', '<')
        line = line.replace('&amp;', '&')
        
        line = line.strip()
        if not line:
            continue
        
        # 添加段落
        if line.startswith('#'):
            # 标题
            level = line.count('#')
            text = line.replace('#', '').strip()
            
            if level == 1:
                p = doc.add_heading(text, level=0)
            elif level == 2:
                p = doc.add_heading(text, level=1)
            elif level == 3:
                p = doc.add_heading(text, level=2)
            else:
                p = doc.add_heading(text, level=3)
        elif line.startswith('•') or line.startswith('-'):
            # 列表
            p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            # 加粗段落
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2].strip())
            run.bold = True
        else:
            # 普通段落
            p = doc.add_paragraph(line)
    
    # 保存文档
    doc.save(docx_file)
    print(f"✅ 转换完成！")
    print(f"📄 输入文件: {md_file}")
    print(f"📄 输出文件: {docx_file}")

if __name__ == '__main__':
    md_file = 'M2.1评测文章_多语言方向_完整版_更新.md'
    docx_file = 'M2.1评测文章_多语言方向_完整版.docx'
    
    if os.path.exists(md_file):
        markdown_to_docx(md_file, docx_file)
    else:
        print(f"❌ 文件不存在: {md_file}")
        print("请确保Markdown文件在同一目录下")
        sys.exit(1)

