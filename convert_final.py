from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import os

def markdown_to_docx(md_file, docx_file):
    """将Markdown文件转换为DOCX"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    html_content = markdown.markdown(md_content, extensions=['extra', 'fenced_code', 'tables'])
    
    doc = Document()
    
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    
    lines = html_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        # 处理代码块
        if '```' in line:
            continue
        
        # 清理HTML标签
        for tag in ['<h1>', '</h1>', '<h2>', '</h2>', '<h3>', '</h3>', 
                    '<p>', '</p>', '<strong>', '</strong>', '<em>', '</em>',
                    '<br>', '<ul>', '</ul>', '<li>', '</li>']:
            line = line.replace(tag, '')
        
        # 处理HTML实体
        entities = {'&gt;': '>', '&lt;': '<', '&amp;': '&', '&quot;': '"'}
        for entity, char in entities.items():
            line = line.replace(entity, char)
        
        line = line.strip()
        if not line:
            continue
        
        # 添加段落
        if line.startswith('#'):
            level = line.count('#')
            text = line.replace('#', '').strip()
            doc.add_heading(text, level=min(level, 3))
        elif line.startswith('•') or line.startswith('-'):
            doc.add_paragraph(line[1:].strip(), style='List Bullet')
        elif '**' in line:
            p = doc.add_paragraph()
            text = line.replace('**', '')
            run = p.add_run(text)
            run.bold = True
        else:
            doc.add_paragraph(line)
    
    doc.save(docx_file)
    print(f"✅ 转换完成: {docx_file}")

if __name__ == '__main__':
    md_file = 'M2.1评测文章_多语言方向_完整版_更新.md'
    docx_file = 'M2.1评测文章_多语言方向_完整版_最终版.docx'
    
    if os.path.exists(md_file):
        markdown_to_docx(md_file, docx_file)

