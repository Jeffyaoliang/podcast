#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
M2.1评测文章Markdown转Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import markdown
import re
import os

def parse_markdown_file(md_file):
    """解析Markdown文件内容"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def extract_code_blocks(content):
    """提取代码块"""
    code_blocks = {}
    pattern = r'```(\w+)?\n([\s\S]*?)```'
    
    def replacer(match):
        lang = match.group(1) or 'text'
        code = match.group(2).strip()
        key = f'__CODE_BLOCK_{len(code_blocks)}__'
        code_blocks[key] = (lang, code)
        return key
    
    content = re.sub(pattern, replacer, content)
    return content, code_blocks

def restore_code_blocks(content, code_blocks):
    """恢复代码块"""
    for key, (lang, code) in code_blocks.items():
        content = content.replace(key, f'```{lang}\n{code}\n```')
    return content

def add_heading_with_style(doc, text, level):
    """添加标题"""
    heading = doc.add_heading('', level=level)
    run = heading.add_run(text)
    run.bold = True
    
    if level == 0:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 102, 204)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 102, 153)
    else:
        run.font.size = Pt(14)

def add_paragraph_with_style(doc, text, style=None):
    """添加段落"""
    para = doc.add_paragraph()
    
    # 处理粗体文本
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
    
    if style == 'indent':
        para.paragraph_format.first_line_indent = Inches(0.5)
    elif style == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_code_block(doc, lang, code):
    """添加代码块"""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 128, 0)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Inches(0.5)

def add_list_item(doc, text):
    """添加列表项"""
    para = doc.add_paragraph()
    para.style = 'List Bullet'
    para.paragraph_format.left_indent = Inches(0.3)
    
    # 处理粗体
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)

def markdown_to_docx(md_file, docx_file):
    """将Markdown转换为DOCX"""
    
    print(f"📖 读取文件: {md_file}")
    
    content = parse_markdown_file(md_file)
    
    # 提取代码块
    content, code_blocks = extract_code_blocks(content)
    
    # 解析Markdown
    md = markdown.Markdown(extensions=['tables', 'fenced_code'])
    html = md.convert(content)
    
    # 创建文档
    doc = Document()
    
    # 设置页面边距
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # 处理HTML
    lines = html.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行和注释
        if not line or line.startswith('<!--'):
            i += 1
            continue
        
        # 跳过代码块标记（已单独处理）
        if line == '```':
            i += 1
            continue
        
        # 清理HTML标签
        line = re.sub(r'<h([1-6])[^>]*>(.*?)</h\1>', lambda m: m.group(2), line)
        line = re.sub(r'<p[^>]*>(.*?)</p>', r'\1', line)
        line = re.sub(r'<strong>(.*?)</strong>', r'\1', line)
        line = re.sub(r'<em>(.*?)</em>', r'\1', line)
        line = re.sub(r'<br\s*/?>', '\n', line)
        line = re.sub(r'<ul[^>]*>', '', line)
        line = re.sub(r'</ul>', '', line)
        line = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1', line)
        line = re.sub(r'<code>(.*?)</code>', r'\1', line)
        line = re.sub(r'<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>', r'\2', line)
        
        # 清理实体
        entities = {
            '&gt;': '>',
            '&lt;': '<',
            '&amp;': '&',
            '&quot;': '"',
            '&nbsp;': ' ',
            '&#39;': "'"
        }
        for entity, char in entities.items():
            line = line.replace(entity, char)
        
        line = line.strip()
        if not line:
            i += 1
            continue
        
        # 判断标题
        if line.startswith('#'):
            hashes = line.count('#')
            text = line.replace('#', '').strip()
            add_heading_with_style(doc, text, min(hashes, 3))
        # 判断列表
        elif line.startswith('•') or line.startswith('- '):
            text = line[1:].strip() if line.startswith('•') else line[1:].strip()
            add_list_item(doc, text)
        # 判断引用
        elif line.startswith('>') or '> ' in line:
            text = line.replace('>', '').strip()
            add_paragraph_with_style(doc, text, 'indent')
        # 其他段落
        else:
            add_paragraph_with_style(doc, line)
        
        i += 1
    
    # 添加代码块到文档
    print("📦 添加代码块...")
    for key, (lang, code) in code_blocks.items():
        if lang:
            heading = doc.add_heading(f'{lang} 代码示例', level=3)
        add_code_block(doc, lang, code)
    
    # 保存文档
    doc.save(docx_file)
    print(f"\n✅ 转换完成！")
    print(f"📄 输出文件: {docx_file}")
    return True

if __name__ == '__main__':
    md_file = 'M2.1评测文章_多语言方向_完整版_更新.md'
    docx_file = 'M2.1评测文章_多语言方向_完整版_最终.docx'
    
    if os.path.exists(md_file):
        markdown_to_docx(md_file, docx_file)
        print(f"\n🎉 成功生成Word文档！")
    else:
        print(f"❌ 文件不存在: {md_file}")

